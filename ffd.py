import docx
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import subprocess
import cups

def set_font_properties(run, font_name, font_size):
    run.font.size = Pt(font_size)
    run.font.name = font_name
    # Apply complex script (CS) font settings for Arabic text
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:cs'), font_name)
    rPr.append(rFonts)
    return run

def add_arabic_text_to_existing_docx(template_path, output_path, names_and_numbers, font_size=72, font_name="Traditional Arabic"):
    doc = docx.Document(template_path)
    majddddtemp  = 0
    # Write names and numbers in Arabic with the specified larger font size and font type
    for name, number in names_and_numbers:
        if majddddtemp == 1 :
            para = doc.add_paragraph()
            run = para.add_run(f"{name}: {number}")
            run = set_font_properties(run, font_name, font_size)
            majddddtemp  = 0
        else:
            para = doc.add_paragraph()
            run = para.add_run(f"{name}: {number}")
            run = set_font_properties(run, font_name, 30)
            majddddtemp  += 1
        # Debug: print out the font size and font name to verify

    
    doc.save(output_path)

def convert_docx_to_pdf(docx_path, pdf_path):
    # Use LibreOffice to convert DOCX to PDF
    command = ['libreoffice', '--headless', '--convert-to', 'pdf', '--outdir', os.path.dirname(pdf_path), docx_path]
    subprocess.run(command, check=True)

def print_pdf(pdf_path, printer_name):
    # Use pycups to print the PDF
    conn = cups.Connection()
    printers = conn.getPrinters()
    
    if printer_name not in printers:
        raise ValueError(f"Printer {printer_name} not found.")
    
    conn.printFile(printer_name, pdf_path, "Document Print", {})

def print_docx(docx_path, printer_name):
    # Convert DOCX to PDF
    pdf_path = docx_path.replace('.docx', '.pdf')
    convert_docx_to_pdf(docx_path, pdf_path)
    
    # Print the PDF
    print_pdf(pdf_path, printer_name)
    
    # Optionally, delete the PDF file after printing
    os.remove(pdf_path)


template_path = "jd.docx"  # Path to the existing DOCX file
output_path = "modified_document.docx"  # Path to save the modified DOCX file
printer_name = "SPRT-"  # Replace with your printer name
from datetime import datetime
# List of names and numbers in Arabic
def run(n, nu ):
    d = datetime.now()
    names_and_numbers = [
(f'{d}' ,''),
        (f"{n}", f"{nu}")
    ]

    # Add Arabic text to the existing DOCX template and increase font size and change font type
    add_arabic_text_to_existing_docx(template_path, output_path, names_and_numbers, font_size=72, font_name="Traditional Arabic")

    # Print the modified DOCX file
    print_docx(output_path, printer_name)



